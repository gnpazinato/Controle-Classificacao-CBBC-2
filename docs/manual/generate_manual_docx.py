"""
Gera a versão editável (Word .docx) do Manual de Uso para a CBBC (v2.3.4).

Reutiliza as capturas renderizadas por app_screens.py (build/*.png) e os
auxiliares de anotação/moldura de generate_manual.py, montando um documento
Word totalmente editável com python-docx.
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

import app_screens  # garante que build/*.png exista (importa rosters etc.)
from screens import load_logo
from generate_manual import annotate, framed_path, flow_diagram, BUILD

OUT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..",
                                   "Manual-Uso-Controle-Classificacao-CBBC.docx"))

BLUE = RGBColor(0x1F, 0x66, 0xB6)
BLUE_DEEP = RGBColor(0x15, 0x4B, 0x82)
ORANGE = RGBColor(0xE8, 0x7B, 0x2B)
TEXT = RGBColor(0x1A, 0x1A, 0x1A)
TEXT2 = RGBColor(0x5A, 0x60, 0x68)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
FONT = "Calibri"

USABLE_IN = 6.5  # largura útil aproximada (A4, margens 2cm) em polegadas


# ---------------------------------------------------------------------------
# Auxiliares de baixo nível
# ---------------------------------------------------------------------------
def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def _set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)


def _no_table_borders(table):
    _tbl_borders(table, "none")


def _tbl_borders(table, val, color="auto", sz="8"):
    """Define w:tblBorders na posição correta do schema (antes de
    shd/tblLayout/tblCellMar/tblLook), evitando que o Word peça reparo."""
    tblPr = table._tbl.tblPr
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), val)
        if val != "none":
            e.set(qn("w:sz"), sz)
            e.set(qn("w:color"), color)
        borders.append(e)
    anchor = None
    for tag in ("w:shd", "w:tblLayout", "w:tblCellMar", "w:tblLook"):
        found = tblPr.find(qn(tag))
        if found is not None:
            anchor = found
            break
    if anchor is not None:
        anchor.addprevious(borders)
    else:
        tblPr.append(borders)


def _runs(paragraph, parts, size=10.5, color=TEXT):
    """parts: lista de (texto, bold). Adiciona runs formatados."""
    for txt, bold in parts:
        r = paragraph.add_run(txt)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.color.rgb = color


def _parse_bold(text):
    """Converte 'a **b** c' em [(a,False),(b,True),(c,False)]."""
    out = []
    for i, chunk in enumerate(text.split("**")):
        if chunk:
            out.append((chunk, i % 2 == 1))
    return out


# ---------------------------------------------------------------------------
# Blocos de alto nível
# ---------------------------------------------------------------------------
def h1(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = BLUE_DEEP
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = BLUE
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def body(doc, text, size=10.5, color=TEXT, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(6)
    for txt, bold in _parse_bold(text):
        r = p.add_run(txt)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = color
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    for txt, bold in _parse_bold(text):
        r = p.add_run(txt)
        r.font.name = FONT
        r.font.size = Pt(10.5)
        r.font.bold = bold
        r.font.color.rgb = TEXT
    return p


def image(doc, path, width_in, caption=None):
    img = Image.open(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(path, width=Emu(int(width_in * 914400)))
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.space_after = Pt(10)
        r = c.add_run(caption)
        r.font.name = FONT
        r.font.size = Pt(9)
        r.font.italic = True
        r.font.color.rgb = TEXT2


def callout(doc, title, items):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _no_table_borders(table)
    cell = table.rows[0].cells[0]
    _shade(cell, "ECF2F9")
    _set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
    cell.paragraphs[0].text = ""
    # título
    tp = cell.paragraphs[0]
    tr = tp.add_run(title)
    tr.font.name = FONT
    tr.font.size = Pt(11)
    tr.font.bold = True
    tr.font.color.rgb = BLUE_DEEP
    tp.paragraph_format.space_after = Pt(4)
    for i, it in enumerate(items, 1):
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        nr = p.add_run(f"{i}  ")
        nr.font.name = FONT
        nr.font.size = Pt(10)
        nr.font.bold = True
        nr.font.color.rgb = BLUE
        for txt, bold in _parse_bold(it):
            r = p.add_run(txt)
            r.font.name = FONT
            r.font.size = Pt(10)
            r.font.bold = bold
            r.font.color.rgb = TEXT
    # borda azul ao redor da caixa
    _tbl_borders(table, "single", color="1F66B6", sz="8")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ---------------------------------------------------------------------------
# Documento
# ---------------------------------------------------------------------------
def build():
    # garante capturas (renderiza as que faltarem)
    render_map = {
        "splash": app_screens.splash, "home": app_screens.home,
        "validation": app_screens.validation_summary, "missing": app_screens.missing_data,
        "setup": app_screens.match_setup,
        "live_normal": lambda: app_screens.live_court("normal"),
        "live_exceeded": lambda: app_screens.live_court("exceeded"),
        "live_bonus": lambda: app_screens.live_court("bonus"),
        "broadcast": app_screens.broadcast_dialog, "viewer": app_screens.public_viewer,
        "live_phone": app_screens.live_phone,
    }
    for name, fn in render_map.items():
        png = os.path.join(BUILD, f"{name}.png")
        if not os.path.exists(png):
            fn().convert("RGB").save(png)
    # logo da capa (fundo transparente -> branco)
    cover_logo = os.path.join(BUILD, "_cover_logo.png")
    if not os.path.exists(cover_logo):
        lg = load_logo(h=600)
        bg = Image.new("RGBA", lg.size, (255, 255, 255, 255))
        bg.alpha_composite(lg)
        bg.convert("RGB").save(cover_logo)

    doc = Document()
    # margens / fonte base
    for sec in doc.sections:
        sec.top_margin = Cm(1.9)
        sec.bottom_margin = Cm(1.7)
        sec.left_margin = Cm(2)
        sec.right_margin = Cm(2)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)

    # rodapé com número de página
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Controle de Classificação CBBC — Manual de Uso (v2.3.4)")
    fr.font.name = FONT
    fr.font.size = Pt(8)
    fr.font.color.rgb = TEXT2

    # ---------------- Capa ----------------
    logo = os.path.join(BUILD, "_cover_logo.png")
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(60)
    cp.add_run().add_picture(logo, width=Cm(6.5))

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(24)
    r = t.add_run("Manual de Uso do Aplicativo")
    r.font.name = FONT
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.color.rgb = BLUE_DEEP

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("Controle de Classificação CBBC")
    r.font.name = FONT
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = BLUE

    for line in ["Conferência da pontuação funcional por equipe",
                 "no basquetebol em cadeira de rodas"]:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.paragraph_format.space_after = Pt(0)
        r = sub.add_run(line)
        r.font.name = FONT
        r.font.size = Pt(12.5)
        r.font.color.rgb = TEXT2

    # caixa de destaque
    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    bcell = box.rows[0].cells[0]
    _shade(bcell, "F8FAFC")
    _set_cell_margins(bcell, 160, 160, 200, 200)
    bp = bcell.paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = bp.add_run("Manual de Uso para a CBBC")
    br.font.name = FONT
    br.font.bold = True
    br.font.size = Pt(11)
    br.font.color.rgb = BLUE_DEEP
    note = bcell.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run("Guia de operação do aplicativo para os classificadores da CBBC durante "
                      "as partidas. Demonstra todas as funcionalidades por meio de uma "
                      "simulação de uso real (Supercopa — IREFES × ADESUL).")
    nr.font.name = FONT
    nr.font.size = Pt(10)
    nr.font.color.rgb = TEXT
    _tbl_borders(box, "single", color="1F66B6", sz="8")

    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ver.paragraph_format.space_before = Pt(36)
    r = ver.add_run("Versão 2.3.4")
    r.font.name = FONT
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = BLUE_DEEP
    org = doc.add_paragraph()
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = org.add_run("Confederação Brasileira de Basquetebol em Cadeira de Rodas")
    r.font.name = FONT
    r.font.size = Pt(9.5)
    r.font.color.rgb = TEXT2
    doc.add_page_break()

    # ---------------- Apresentação ----------------
    h1(doc, "Apresentação")
    body(doc, "O **Controle de Classificação CBBC** é um aplicativo desenvolvido para os "
              "**classificadores da CBBC** acompanharem, em tempo real, a soma dos "
              "**pontos funcionais** das equipes durante as partidas de basquetebol em "
              "cadeira de rodas. A cada atleta em quadra, o app soma automaticamente as classes "
              "funcionais e avisa imediatamente quando o limite permitido é ultrapassado — "
              "eliminando erros de conta manual e dando mais segurança e agilidade ao trabalho.")
    body(doc, "Além do controle, a versão 2.3.4 introduz a **transmissão ao vivo da quadra**: "
              "com um toque, o app gera um **QR Code e um link público** que qualquer pessoa "
              "abre no navegador para ver o posicionamento e o placar de pontos em tempo quase "
              "real — ideal para a mesa e o público.")
    h2(doc, "Conteúdo")
    for it in ["1. Visão geral e fluxo de uso",
               "2. Abertura e tela inicial",
               "3. Carregamento e conferência dos atletas (planilha .xlsx ou PDF)",
               "4. Configuração da partida (equipes, camisetas, pontuação, bonificações)",
               "5. Quadra ao vivo — operação durante o jogo",
               "6. Transmissão ao vivo com QR Code e página pública",
               "7. Regras de pontuação e bonificação",
               "8. Boas práticas, requisitos e privacidade"]:
        bullet(doc, it)
    body(doc, "Os registros de tela deste manual reproduzem uma **simulação de uso real** da "
              "Supercopa, com as equipes **IREFES** e **ADESUL** a partir da planilha de atletas.",
         size=9.5, color=TEXT2)
    doc.add_page_break()

    # ---------------- 1. Visão geral ----------------
    h1(doc, "1. Visão geral e fluxo de uso")
    body(doc, "O uso do aplicativo segue cinco etapas simples, sempre na mesma ordem. Os dados da "
              "partida funcionam **100% offline**; a internet só é necessária para carregar as "
              "**fotos** das atletas e para a **transmissão ao vivo** — recursos opcionais.")
    image(doc, flow_diagram(), USABLE_IN, "Fluxo completo: da planilha à transmissão pública da quadra.")
    h2(doc, "Principais benefícios")
    for it in ["**Soma automática** dos pontos funcionais das cinco atletas em quadra, por equipe.",
               "**Alerta visual e por vibração** assim que o limite de pontos é excedido.",
               "**Bonificações** (Sub-16, Sub-23 e atleta feminina) aplicadas automaticamente, "
               "elevando o teto para 15,0 pontos quando há atleta da categoria em quadra.",
               "**Identificação por foto**, número de camisa e classe de cada atleta na quadra.",
               "**Funciona offline** — não depende de internet para conferir a partida.",
               "**Transmissão pública** da quadra por QR Code / link, sem instalar nada no "
               "dispositivo do espectador.",
               "Aceita **planilha .xlsx ou PDF** e mantém a sessão salva caso o app seja fechado."]:
        bullet(doc, it)
    doc.add_page_break()

    # ---------------- 2. Abertura ----------------
    h1(doc, "2. Abertura e tela inicial")
    body(doc, "Ao abrir, o app exibe a identidade visual da CBBC e segue para a tela inicial, onde "
              "a partida é iniciada e os modelos de planilha podem ser baixados.")
    # duas imagens lado a lado
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _no_table_borders(tbl)
    for cell, path in zip(tbl.rows[0].cells,
                          [framed_path("splash.png"),
                           annotate("home.png", [(95, 595), (330, 978), (650, 978), (240, 1414)])]):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Cm(6.4))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run("À esquerda, a tela de abertura. À direita, a tela inicial.")
    r.font.name = FONT
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = TEXT2
    callout(doc, "Tela inicial — passo a passo", [
        "**Carregar planilha (.xlsx) ou PDF** — abre o seletor de arquivos para importar a "
        "lista de atletas e iniciar a conferência.",
        "**Baixar modelo “Aba única”** — planilha-modelo com todos os clubes em uma só aba.",
        "**Baixar modelo “Por clube”** — planilha-modelo com uma aba para cada clube.",
        "**Versão do app e aviso “Dados offline”** — confirmam que a partida funciona sem internet.",
    ])
    doc.add_page_break()

    # ---------------- 3. Carregamento ----------------
    h1(doc, "3. Carregamento e conferência dos atletas")
    body(doc, "A planilha (ou PDF) deve conter as colunas **clube, classe, atleta, camisa, data de "
              "nascimento, gênero** e, opcionalmente, **foto** (link público). O app reconhece "
              "cabeçalhos em português e inglês. Após o carregamento, a tela de **Resumo da importação** "
              "mostra tudo o que foi lido e permite ajustar qualquer dado antes do jogo.")
    image(doc, annotate("validation.png", [(75, 250), (688, 250), (560, 558), (390, 688), (415, 1561)]),
          3.1, "Resumo da importação — Supercopa: 4 clubes e 42 atletas reconhecidos.")
    callout(doc, "Resumo da importação", [
        "**Contadores** de clubes e atletas encontrados no arquivo.",
        "**Selo de status** — verde quando o arquivo está pronto; vermelho quando há erros.",
        "**Renomear / excluir** equipe diretamente na lista.",
        "**Edição rápida** de cada atleta: camisa, nome, foto, nascimento, gênero e classe.",
        "**Continuar** — avança para a configuração da partida (habilita só sem erros).",
    ])
    h2(doc, "Quando há informações faltando")
    body(doc, "Se algum dado obrigatório estiver ausente (por exemplo, classe funcional ou data de "
              "nascimento), o app destaca o problema e indica exatamente qual atleta e linha corrigir, "
              "evitando que a partida comece com dados incompletos.")
    image(doc, framed_path("missing.png"), 3.0,
          "Tela “Dados pendentes”: erros agrupados por tipo, com clube, atleta e linha.")
    doc.add_page_break()

    # ---------------- 4. Configuração ----------------
    h1(doc, "4. Configuração da partida")
    body(doc, "Nesta tela definem-se as duas equipes do confronto, as cores das camisetas, a pontuação "
              "máxima e as bonificações válidas para a competição. Na simulação, configuramos "
              "**IREFES (Equipe A)** contra **ADESUL (Equipe B)**.")
    image(doc, annotate("setup.png", [(560, 216), (640, 331), (600, 704), (620, 792), (560, 1000), (600, 1252)]),
          3.0, "Configuração da partida IREFES × ADESUL, com bonificação Sub-23 ativada.")
    callout(doc, "Configuração da partida", [
        "**Equipe A e Equipe B** — selecionadas entre os clubes importados.",
        "**Cor da camiseta** de cada equipe — usada nos ícones e chips em quadra.",
        "**Pontuação máxima por equipe** — limite padrão de 14,0 (ajustável de 7,0 a 16,0).",
        "**Bonificações da competição** — Sub-16, Sub-23 e atleta feminina.",
        "**Sub-23 ativada** — atletas da categoria recebem destaque e elevam o teto para 15,0.",
        "**Datas de referência** — conferidas com o tablet; usadas no cálculo das idades.",
    ])
    doc.add_page_break()

    # ---------------- 5. Quadra ao vivo ----------------
    h1(doc, "5. Quadra ao vivo — operação durante o jogo")
    body(doc, "É a tela principal, usada durante a partida. Em tablet, mostra as duas listas de atletas "
              "nas laterais e a **quadra** ao centro. Toca-se no nome (ou na foto em quadra) para "
              "colocar ou retirar uma atleta; o app soma os pontos e atualiza o placar instantaneamente.")
    image(doc, annotate("live_normal.png", [(2150, 70), (2230, 70), (690, 168), (180, 270), (770, 800)],
                        frame="tablet"), USABLE_IN,
          "Quadra ao vivo (tablet) — IREFES 13,5 e ADESUL 13,5, ambas dentro do limite de 14,0.")
    callout(doc, "Quadra ao vivo", [
        "**Botão de transmissão** — inicia/abre a transmissão pública (seção 6).",
        "**Ajuste de pontuação máxima** — altera o limite sem sair da partida.",
        "**Placar por equipe** — soma atual / limite; fica vermelho ao exceder.",
        "**Listas das equipes** — toque para colocar/retirar atletas (máx. 5 por equipe).",
        "**Quadra** — mostra as atletas posicionadas, com foto, número e classe; "
        "abaixo, o seletor de estilo de quadra (clara/escura).",
    ])
    h2(doc, "Alerta de limite excedido")
    body(doc, "Quando a soma das cinco atletas ultrapassa o limite, o placar daquela equipe fica "
              "**vermelho** com a mensagem **“Limite excedido.”** e o tablet **vibra** — "
              "o erro é percebido na hora, mesmo sem olhar a tela.")
    image(doc, framed_path("live_exceeded.png", frame="tablet"), USABLE_IN,
          "IREFES com 15,0 pontos para um limite de 14,0: placar vermelho e aviso de limite excedido.")
    doc.add_page_break()

    h2(doc, "Bonificação em quadra")
    body(doc, "Com uma bonificação ativa (aqui, Sub-23), as atletas da categoria aparecem com uma "
              "**estrela laranja** na lista e na quadra. Enquanto houver ao menos uma delas em "
              "quadra, o **teto da equipe sobe para 15,0** pontos sem gerar alerta — o placar mostra "
              "“/ 15.0” com a estrela.")
    image(doc, framed_path("live_bonus.png", frame="tablet"), USABLE_IN,
          "Bonificação Sub-23 em quadra: teto elevado para 15,0 e estrela ao lado das atletas da categoria.")
    h2(doc, "Operação no celular")
    body(doc, "Em telas menores, a quadra é organizada em três abas — **Equipe A**, **Quadra** e "
              "**Equipe B** — preservando todas as funções.")
    image(doc, framed_path("live_phone.png"), 3.2,
          "Mesma partida em um celular: navegação por abas, com o placar sempre visível no topo.")
    doc.add_page_break()

    # ---------------- 6. Transmissão ----------------
    h1(doc, "6. Transmissão ao vivo com QR Code")
    body(doc, "Novidade da versão 2.3, a transmissão pública permite **espelhar a quadra** para "
              "qualquer pessoa, sem instalar nada. Ao tocar no **ícone de transmissão** na barra "
              "superior, o app cria a sessão e exibe uma janela com **QR Code** e **link** "
              "(no formato …/v/<código>). Basta apontar a câmera ou copiar o link.")
    image(doc, framed_path("broadcast.png", frame="tablet"), USABLE_IN,
          "Janela de transmissão: QR Code e link público copiável, sobre a partida em andamento.")
    callout(doc, "Como transmitir", [
        "Toque no **ícone de transmissão** (canto superior direito). Ele fica **laranja** "
        "enquanto a transmissão está no ar.",
        "Mostre o **QR Code** para o público ou **copie o link** para enviar/compartilhar.",
        "A página pública atualiza em **~1 segundo** e aceita **vários espectadores** ao mesmo tempo.",
        "Toque em **Encerrar** para finalizar; a sessão também expira sozinha após 1 hora de inatividade.",
    ])
    body(doc, "A transmissão funciona **apenas com internet** — exatamente como as fotos das atletas. "
              "Sem conexão, a partida segue 100% normal e o botão apenas avisa que é preciso estar online.",
         size=9.5, color=TEXT2)
    doc.add_page_break()

    h2(doc, "A página pública (no navegador)")
    body(doc, "Quem abre o link vê **somente a quadra** — chips das atletas e o placar nos cantos — "
              "sobre fundo preto, sem botões nem menus. É a página ideal para projetar ou "
              "compartilhar com o público.")
    image(doc, framed_path("viewer.png", frame="browser"), 4.2,
          "Página pública /v/<código> aberta no navegador: a mesma quadra ao vivo, pronta para exibição.")
    callout(doc, "Página pública", [
        "Abre em **qualquer navegador** (celular, computador ou TV) — não precisa do app.",
        "Mostra **posicionamento, fotos, números, classes** e o **placar de pontos** de cada equipe.",
        "Reflete **em tempo quase real** cada toque feito pelo classificador no tablet.",
        "Não expõe dados sensíveis nem controles — é somente visualização.",
    ])
    doc.add_page_break()

    # ---------------- 7. Regras ----------------
    h1(doc, "7. Regras de pontuação e bonificação")
    body(doc, "O app segue o padrão de classes funcionais adotado pela CBBC. Cada atleta possui uma "
              "classe de **1,0 a 4,5** (em incrementos de 0,5). A pontuação da equipe é a soma das "
              "classes das **cinco atletas em quadra**.")
    rules = [("Conceito", "Como funciona no app"),
             ("Classes funcionais", "1,0 · 1,5 · 2,0 · 2,5 · 3,0 · 3,5 · 4,0 · 4,5"),
             ("Atletas em quadra", "Até 5 por equipe; a soma das classes é o total da equipe."),
             ("Limite padrão", "14,0 pontos (configurável de 7,0 a 16,0)."),
             ("Limite excedido", "Placar vermelho, mensagem de alerta e vibração do tablet."),
             ("Bonificação", "Sub-16, Sub-23 e/ou atleta feminina, definidas na configuração."),
             ("Teto com bonificação", "15,0 pontos enquanto houver atleta bonificada em quadra."),
             ("Cálculo da idade", "Pela data de término da competição (datas de referência).")]
    table = doc.add_table(rows=len(rules), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(5.2)
    table.columns[1].width = Cm(11.3)
    for ri, (a, b) in enumerate(rules):
        c0, c1 = table.rows[ri].cells
        c0.width = Cm(5.2)
        c1.width = Cm(11.3)
        for cell, txt, isval in ((c0, a, False), (c1, b, False)):
            cell.paragraphs[0].text = ""
            run = cell.paragraphs[0].add_run(txt)
            run.font.name = FONT
            run.font.size = Pt(9.6)
            if ri == 0:
                run.font.bold = True
                run.font.color.rgb = WHITE
            else:
                run.font.color.rgb = BLUE_DEEP if cell is c0 else TEXT
                if cell is c0:
                    run.font.bold = True
        if ri == 0:
            _shade(c0, "1F66B6")
            _shade(c1, "1F66B6")
        elif ri % 2 == 0:
            _shade(c0, "F8FAFC")
            _shade(c1, "F8FAFC")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    h2(doc, "Sobre as bonificações")
    for it in ["**Sub-16**: vale enquanto a atleta não completa 17 anos até o término da competição.",
               "**Sub-23**: vale enquanto a atleta não completa 24 anos até o término da competição.",
               "**Atleta feminina**: para competições masculinas que concedem bonificação a atletas femininas.",
               "A estrela laranja (na lista e na quadra) indica, durante o jogo, quais atletas se "
               "enquadram na bonificação ativa."]:
        bullet(doc, it)
    doc.add_page_break()

    # ---------------- 8. Boas práticas ----------------
    h1(doc, "8. Boas práticas, requisitos e privacidade")
    h2(doc, "Recomendações para o uso em quadra")
    for it in ["Prepare a planilha com antecedência e confira o **Resumo da importação** antes do jogo.",
               "Garanta que todas as atletas tenham **classe funcional** preenchida — é o dado que soma os pontos.",
               "Para usar fotos e transmissão, conecte o tablet à **internet** antes da partida.",
               "Confirme as **datas de referência** para que as bonificações por idade sejam calculadas corretamente.",
               "Mantenha o tablet carregado: a tela permanece ligada durante a partida."]:
        bullet(doc, it)
    h2(doc, "Requisitos")
    for it in ["**Tablet ou celular Android** (recomendado tablet para a visão completa da quadra).",
               "**Conferência da partida**: funciona totalmente offline.",
               "**Fotos e transmissão ao vivo**: exigem conexão com a internet.",
               "Arquivo de atletas em **.xlsx** ou **PDF** com texto selecionável."]:
        bullet(doc, it)
    h2(doc, "Privacidade e segurança")
    for it in ["Os dados da partida ficam **no próprio dispositivo** (sessão salva localmente).",
               "A transmissão publica **apenas a quadra** (posições, números, classes e placar) — "
               "sem expor a planilha ou dados pessoais sensíveis.",
               "Cada transmissão usa um **código aleatório** e **expira sozinha** após 1 hora de "
               "inatividade; pode ser encerrada a qualquer momento."]:
        bullet(doc, it)
    doc.add_paragraph()
    body(doc, "Este manual acompanha a versão **2.3.4** do aplicativo Controle de Classificação CBBC. "
              "As telas reproduzem fielmente a interface do app em uma simulação real da Supercopa "
              "(IREFES × ADESUL). Em dispositivos conectados à internet, os retratos azuis das atletas "
              "são substituídos pelas **fotos reais** carregadas dos links da planilha.",
         size=9.5, color=TEXT2)
    sign = doc.add_paragraph()
    sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sign.add_run("Confederação Brasileira de Basquetebol em Cadeira de Rodas (CBBC)")
    r.font.name = FONT
    r.font.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = BLUE_DEEP

    doc.save(OUT)
    print("DOCX gerado:", OUT, os.path.getsize(OUT), "bytes")


if __name__ == "__main__":
    build()
